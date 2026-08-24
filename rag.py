import json
import threading
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any

import faiss
try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None

try:
    from rapidocr_onnxruntime import RapidOCR
    _ocr_engine = RapidOCR()
except Exception:
    _ocr_engine = None

try:
    import docx
except ImportError:
    docx = None

import numpy as np
import ollama

from config import (
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    TOP_K,
    get_user_notes_dir,
    get_user_vector_dir
)
from database import (
    create_document_job,
    update_job_status,
    save_document_extracted_text,
    delete_user_document
)


import re
import zipfile

# --------------------------------------------------
# MULTI-FORMAT TEXT EXTRACTION & OCR
# --------------------------------------------------

def extract_image_text(image_path: str) -> str:
    """Extracts text from an image file (PNG, JPG, JPEG, WEBP) using RapidOCR with PyMuPDF fallback."""
    extracted_lines = []
    
    # 1. Primary: RapidOCR engine
    if _ocr_engine is not None:
        try:
            result, _ = _ocr_engine(str(image_path))
            if result:
                for line in result:
                    if line and len(line) > 1 and line[1]:
                        extracted_lines.append(str(line[1]).strip())
                text = "\n".join(l for l in extracted_lines if l).strip()
                if text:
                    return text
        except Exception:
            pass

    # 2. Secondary fallback: PyMuPDF pixmap extraction if available
    if fitz is not None:
        try:
            doc = fitz.open(str(image_path))
            text = "".join(page.get_text() for page in doc).strip()
            doc.close()
            if text:
                return text
        except Exception:
            pass

    return ""


def extract_docx_text(docx_path: str) -> str:
    """
    Extracts structured text from a .docx Word document.
    Extracts paragraphs, headings, table cells, headers/footers, and text box content.
    """
    full_text = []

    # 1. Primary python-docx parser
    if docx is not None:
        try:
            doc = docx.Document(str(docx_path))
            
            # Paragraphs & Headings
            for para in doc.paragraphs:
                p_text = para.text.strip()
                if p_text:
                    full_text.append(p_text)
                    
            # Tables (extract rows and cells with structured pipes)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate repeated adjacent cells in merged rows
                    unique_cells = []
                    for c in row_cells:
                        if not unique_cells or unique_cells[-1] != c:
                            unique_cells.append(c)
                    if unique_cells:
                        full_text.append(" | ".join(unique_cells))

            # Headers & Footers
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for hp in section.header.paragraphs:
                        if hp.text.strip() and hp.text.strip() not in full_text:
                            full_text.insert(0, hp.text.strip())
                if section.footer and section.footer.paragraphs:
                    for fp in section.footer.paragraphs:
                        if fp.text.strip() and fp.text.strip() not in full_text:
                            full_text.append(fp.text.strip())

            # Text boxes in drawing shapes (XML xpath fallback)
            try:
                txbx_nodes = doc._element.xpath('.//w:txbxContent//w:p')
                for node in txbx_nodes:
                    node_text = "".join(node.itertext()).strip()
                    if node_text and node_text not in full_text:
                        full_text.append(node_text)
            except Exception:
                pass

            result = "\n\n".join(full_text).strip()
            if result:
                return result
        except Exception:
            pass

    # 2. Fallback: Direct zip XML extraction from word/document.xml
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml').decode('utf-8', errors='ignore')
                # Extract text inside <w:t> tags
                text_runs = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml_content)
                if text_runs:
                    clean_text = " ".join(text_runs).strip()
                    if clean_text:
                        return clean_text
    except Exception:
        pass

    return ""


def extract_doc_text(doc_path: str) -> str:
    """
    Extracts text from legacy Microsoft Word .doc files, RTF files, or XML/HTML Word files.
    """
    try:
        with open(doc_path, "rb") as f:
            raw_bytes = f.read()

        if not raw_bytes:
            return ""

        # Case 1: Check for RTF signature (\rtf)
        if raw_bytes.startswith(b'{\\rtf'):
            text_str = raw_bytes.decode('latin-1', errors='ignore')
            # Remove RTF control words and groups
            text_clean = re.sub(r'\\[a-zA-Z0-9\-]+ ?', ' ', text_str)
            text_clean = re.sub(r'[{}]', ' ', text_clean)
            lines = [l.strip() for l in text_clean.splitlines() if l.strip()]
            if lines:
                return "\n".join(lines)

        # Case 2: Check for ZIP signature PK\x03\x04 (renamed .docx)
        if raw_bytes.startswith(b'PK\x03\x04'):
            docx_res = extract_docx_text(doc_path)
            if docx_res:
                return docx_res

        # Case 3: Check for Word XML/HTML format (<?xml or <html)
        if raw_bytes.strip().startswith(b'<?xml') or b'<html' in raw_bytes[:200].lower():
            text_str = raw_bytes.decode('utf-8', errors='ignore')
            text_clean = re.sub(r'<[^>]+>', ' ', text_str)
            text_clean = re.sub(r'\s+', ' ', text_clean).strip()
            if len(text_clean) > 20:
                return text_clean

        # Case 4: OLE2 Binary Compound File (.doc format)
        # Extract ASCII and UTF-16LE text streams from binary content
        ascii_runs = re.findall(r'[\x20-\x7E\t\r\n]{4,}', raw_bytes.decode('latin-1', errors='ignore'))
        
        # Also extract UTF-16LE strings common in Word binary streams
        utf16_runs = []
        try:
            raw_u16 = raw_bytes.decode('utf-16-le', errors='ignore')
            utf16_runs = re.findall(r'[\x20-\x7E\t\r\n]{4,}', raw_u16)
        except Exception:
            pass

        # Combine, filter out binary junk and known OLE metadata keywords
        combined = []
        skip_patterns = {'Root Entry', 'WordDocument', '1Table', '0Table', 'Data', 'SummaryInformation', 'DocumentSummaryInformation', 'CompObj'}
        
        for run in ascii_runs + utf16_runs:
            run_clean = run.strip()
            if len(run_clean) >= 4 and run_clean not in skip_patterns:
                # Check ratio of alphabetic characters to avoid binary noise
                alpha_count = sum(1 for c in run_clean if c.isalnum() or c.isspace())
                if alpha_count / max(1, len(run_clean)) > 0.65:
                    combined.append(run_clean)

        # Deduplicate while preserving order
        seen = set()
        deduped = []
        for c in combined:
            if c not in seen:
                seen.add(c)
                deduped.append(c)

        if deduped:
            return "\n".join(deduped)

    except Exception:
        pass

    return ""


def extract_plain_text(file_path: str) -> str:
    """Extracts text from .txt, .md, or plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        try:
            with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""


def extract_pdf_pages(pdf_path: str, force_ocr: bool = False, progress_callback: Optional[Any] = None) -> List[Dict[str, Any]]:
    """Extracts text and page numbers from a PDF using PyMuPDF and RapidOCR for scanned pages."""
    if fitz is None:
        return []
    document = fitz.open(str(pdf_path))
    pages = []
    source_name = Path(pdf_path).name
    total_pages = document.page_count

    for page_number, page in enumerate(document, start=1):
        text = page.get_text().strip()

        # If page has no selectable text (scanned PDF) or force_ocr is True, run OCR on rendered page pixmap
        if (len(text) < 15 or force_ocr) and _ocr_engine is not None:
            try:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                result, _ = _ocr_engine(img_bytes)
                if result:
                    lines = [line[1] for line in result if line and len(line) > 1 and line[1]]
                    ocr_text = "\n".join(lines).strip()
                    if force_ocr:
                        text = text + "\n" + ocr_text
                    elif len(ocr_text) > len(text):
                        text = ocr_text
            except Exception:
                pass

        if text:
            pages.append({
                "text": text,
                "page": page_number,
                "source": source_name
            })
            
        if progress_callback:
            try:
                progress_callback(page_number, total_pages)
            except Exception:
                pass

    document.close()
    return pages


def extract_document_pages(file_path: str, force_ocr: bool = False, progress_callback: Optional[Any] = None) -> List[Dict[str, Any]]:
    """
    Routes the file to the appropriate extractor based on its extension.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    source_name = path.name

    if ext == ".pdf":
        return extract_pdf_pages(file_path, force_ocr=force_ocr, progress_callback=progress_callback)
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tiff"):
        text = extract_image_text(file_path)
        if text:
            return [{"text": text, "page": 1, "source": source_name}]
        return []
    elif ext == ".docx":
        text = extract_docx_text(file_path)
        if text:
            return [{"text": text, "page": 1, "source": source_name}]
        return []
    elif ext == ".doc":
        text = extract_doc_text(file_path)
        if text:
            return [{"text": text, "page": 1, "source": source_name}]
        return []
    elif ext in (".txt", ".md"):
        text = extract_plain_text(file_path)
        if text:
            return [{"text": text, "page": 1, "source": source_name}]
        return []
    else:
        text = extract_plain_text(file_path)
        if text:
            return [{"text": text, "page": 1, "source": source_name}]
        return []


def get_full_document_text(file_path: str) -> str:
    """Returns all extracted raw text from any supported document format."""
    pages = extract_document_pages(file_path)
    return "\n\n".join(p["text"] for p in pages if p.get("text"))


# --------------------------------------------------
# CHUNKING & EMBEDDINGS
# --------------------------------------------------

def chunk_text(text: str) -> List[str]:
    """Splits raw text into overlapping token/character chunks."""
    text = text.replace("\n", " ").strip()
    chunks = []
    start = 0
    step = max(1, CHUNK_SIZE - CHUNK_OVERLAP)

    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start += step

    return chunks


def create_chunks(
    file_path: str,
    subject: str,
    user_id: int = 1,
    document_id: Optional[int] = None
) -> List[Dict[str, Any]]:
    """Extracts text and generates metadata-rich chunks for RAG from any supported file."""
    pages = extract_document_pages(file_path)
    chunks = []
    source_name = Path(file_path).name

    for page in pages:
        page_chunks = chunk_text(page["text"])
        for chunk in page_chunks:
            chunks.append({
                "text": chunk,
                "subject": subject,
                "page": page.get("page", 1),
                "source": source_name,
                "user_id": user_id,
                "document_id": document_id
            })

    return chunks


_cached_embedding_model: Optional[str] = None


def get_active_embedding_model() -> str:
    """Returns cached or auto-detected available Ollama embedding model."""
    global _cached_embedding_model
    if _cached_embedding_model:
        return _cached_embedding_model

    try:
        model_list = ollama.list()
        available = [m.model for m in getattr(model_list, "models", [])]
        if EMBEDDING_MODEL in available:
            _cached_embedding_model = EMBEDDING_MODEL
            return _cached_embedding_model

        for candidate in available:
            if any(k in candidate.lower() for k in ("embed", "minilm", "nomic", "bge", "qwen")):
                _cached_embedding_model = candidate
                return _cached_embedding_model

        if available:
            _cached_embedding_model = available[0]
            return _cached_embedding_model
    except Exception:
        pass

    _cached_embedding_model = EMBEDDING_MODEL
    return _cached_embedding_model


def get_embeddings(
    texts: List[str],
    batch_size: int = 4,
    progress_callback: Optional[Any] = None
) -> np.ndarray:
    """Generates normalized vector embeddings via Ollama embedding model with batching and progress tracking."""
    if not texts:
        return np.empty((0, 0), dtype="float32")

    model_name = get_active_embedding_model()
    all_embeddings: List[List[float]] = []
    total = len(texts)

    for i in range(0, total, batch_size):
        batch = texts[i:i + batch_size]
        try:
            response = ollama.embed(
                model=model_name,
                input=batch
            )
            raw_emb = response.get("embeddings", [])
            if isinstance(raw_emb, list) and len(raw_emb) > 0:
                if isinstance(raw_emb[0], list):
                    all_embeddings.extend(raw_emb)
                else:
                    all_embeddings.append(raw_emb)
        except Exception:
            # Fallback to single text embedding if batch fails
            for single_text in batch:
                single_resp = ollama.embed(
                    model=model_name,
                    input=single_text
                )
                raw_s = single_resp.get("embeddings", [])
                if isinstance(raw_s, list) and len(raw_s) > 0:
                    if isinstance(raw_s[0], list):
                        all_embeddings.extend(raw_s)
                    else:
                        all_embeddings.append(raw_s)

        if progress_callback:
            try:
                progress_callback(min(total, i + len(batch)), total)
            except Exception:
                pass

    embeddings = np.array(
        all_embeddings,
        dtype="float32"
    )

    faiss.normalize_L2(embeddings)
    return embeddings


def build_index(chunks: List[Dict[str, Any]]) -> Optional[faiss.IndexFlatIP]:
    """Builds a normalized FAISS cosine-similarity index from chunk embeddings."""
    if not chunks:
        return None

    texts = [item["text"] for item in chunks]
    embeddings = get_embeddings(texts)

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)
    index.add(embeddings)
    return index


def search(
    question: str,
    chunks: List[Dict[str, Any]],
    index: Optional[faiss.IndexFlatIP],
    subject: Optional[str] = None,
    top_k: int = TOP_K,
    user_id: Optional[int] = None
) -> List[Dict[str, Any]]:
    """Performs isolated vector search for the user and subject."""
    if index is None or not chunks:
        return []

    query_embedding = get_embeddings([question])
    k_val = min(top_k * 3, len(chunks))
    scores, indices = index.search(query_embedding, k_val)

    results = []
    for score, index_number in zip(scores[0], indices[0]):
        if index_number < 0 or index_number >= len(chunks):
            continue

        item = chunks[index_number]

        # Enforce subject filter if specified
        if subject and item.get("subject") != subject:
            continue

        # Enforce user isolation if user_id specified
        if user_id is not None and item.get("user_id") is not None and item.get("user_id") != user_id:
            continue

        results.append({
            **item,
            "score": float(score)
        })

        if len(results) >= top_k:
            break

    return results


def save_index(index, chunks: List[Dict[str, Any]], index_path: Path, metadata_path: Path):
    """Saves FAISS index and JSON metadata persistently."""
    if index is None:
        return

    index_path = Path(index_path)
    metadata_path = Path(metadata_path)

    index_path.parent.mkdir(parents=True, exist_ok=True)
    metadata_path.parent.mkdir(parents=True, exist_ok=True)

    faiss.write_index(index, str(index_path))

    with open(metadata_path, "w", encoding="utf-8") as file:
        json.dump(chunks, file, ensure_ascii=False, indent=2)


def load_index(index_path: Path, metadata_path: Path) -> Tuple[Optional[faiss.IndexFlatIP], List[Dict[str, Any]]]:
    """Loads a FAISS index and metadata chunks from disk."""
    index_path = Path(index_path)
    metadata_path = Path(metadata_path)

    if not index_path.exists() or not metadata_path.exists():
        return None, []

    index = faiss.read_index(str(index_path))
    with open(metadata_path, "r", encoding="utf-8") as file:
        chunks = json.load(file)

    return index, chunks


def get_user_subject_index_paths(user_id: int, subject: str) -> Tuple[Path, Path]:
    """Returns the user-isolated index and metadata file paths."""
    user_vec_dir = get_user_vector_dir(user_id)
    index_path = user_vec_dir / f"{subject}.index"
    metadata_path = user_vec_dir / f"{subject}.json"
    return index_path, metadata_path


def load_user_subject_index(user_id: int, subject: str) -> Tuple[Optional[faiss.IndexFlatIP], List[Dict[str, Any]]]:
    """Loads vector index and metadata for a specific user and subject."""
    if not subject or subject == "General":
        return None, []
    index_path, metadata_path = get_user_subject_index_paths(user_id, subject)
    return load_index(index_path, metadata_path)


def search_user_notes(
    question: str,
    user_id: int,
    document_name: Optional[str] = None,
    subject: Optional[str] = None,
    top_k: int = TOP_K
) -> List[Dict[str, Any]]:
    """
    Directly and persistently searches vector stores isolated for a specific user.
    Can search across all completed notes, or filter to a specific subject or document.
    """
    user_vec_dir = get_user_vector_dir(user_id)
    if not user_vec_dir.exists():
        return []

    # Find index files to query
    if subject and subject not in ("General", "All my completed notes"):
        index_files = [user_vec_dir / f"{subject}.index"]
    else:
        index_files = list(user_vec_dir.glob("*.index"))

    all_results = []
    for idx_file in index_files:
        meta_file = idx_file.with_suffix(".json")
        if not idx_file.exists() or not meta_file.exists():
            continue
        try:
            index, chunks = load_index(idx_file, meta_file)
            if index is None or not chunks:
                continue

            res = search(
                question=question,
                chunks=chunks,
                index=index,
                subject=None,
                top_k=top_k * 3,
                user_id=user_id
            )
            for item in res:
                if document_name and document_name != "All my completed notes":
                    if item.get("source") != document_name and item.get("filename") != document_name:
                        continue
                all_results.append(item)
        except Exception:
            continue

    all_results.sort(key=lambda x: x.get("score", 0.0), reverse=True)
    return all_results[:top_k]


def _filter_and_rebuild_index(
    index: Optional[faiss.IndexFlatIP],
    chunks: List[Dict[str, Any]],
    document_id: int,
    filename: str
) -> Tuple[Optional[faiss.IndexFlatIP], List[Dict[str, Any]]]:
    """
    Instantly extracts existing vectors from FAISS index and rebuilds without re-running Ollama embeddings.
    Falls back to computing embeddings if index vectors cannot be extracted.
    """
    if not chunks:
        return None, []

    keep_indices = []
    remaining_chunks = []
    for idx, c in enumerate(chunks):
        if c.get("document_id") != document_id and c.get("source") != filename and c.get("filename") != filename:
            keep_indices.append(idx)
            remaining_chunks.append(c)

    if not remaining_chunks:
        return None, []

    # 1. Fast path: Extract vectors instantly from existing FAISS index (0ms)
    if index is not None and index.ntotal == len(chunks):
        try:
            all_vecs = index.reconstruct_n(0, index.ntotal)
            remaining_vecs = all_vecs[keep_indices]
            new_idx = faiss.IndexFlatIP(index.d)
            new_idx.add(remaining_vecs)
            return new_idx, remaining_chunks
        except Exception:
            try:
                vec_list = [index.reconstruct(i) for i in keep_indices]
                remaining_vecs = np.array(vec_list, dtype="float32")
                new_idx = faiss.IndexFlatIP(index.d)
                new_idx.add(remaining_vecs)
                return new_idx, remaining_chunks
            except Exception:
                pass

    # 2. Fallback path if index was out of sync or missing
    try:
        texts = [c["text"] for c in remaining_chunks]
        embeddings = get_embeddings(texts)
        if len(embeddings) > 0 and embeddings.shape[0] > 0:
            new_idx = faiss.IndexFlatIP(embeddings.shape[1])
            new_idx.add(embeddings)
            return new_idx, remaining_chunks
    except Exception:
        pass

    return None, remaining_chunks


def delete_document_data(
    user_id: int,
    document_id: int,
    filename: str,
    subject: str,
    file_path: Optional[str] = None
) -> Tuple[bool, str]:
    """
    Instantly and completely purges an uploaded document:
    1. Removes the physical file from disk.
    2. Filters vector index & chunks using existing FAISS vectors in memory (instant, no Ollama calls).
    3. Deletes document and job records from database.
    """
    try:
        # 1. Remove physical file from disk
        candidate_paths = []
        if file_path:
            candidate_paths.append(Path(file_path))
        user_notes_dir = get_user_notes_dir(user_id, subject)
        candidate_paths.append(user_notes_dir / filename)

        for p in candidate_paths:
            try:
                if p.exists() and p.is_file():
                    p.unlink()
            except Exception:
                pass

        # 2. Update Vector Store & Chunks for the primary subject (instant)
        if subject and subject != "General":
            index_path, metadata_path = get_user_subject_index_paths(user_id, subject)
            if index_path.exists() and metadata_path.exists():
                try:
                    index, existing_chunks = load_index(index_path, metadata_path)
                    new_idx, remaining_chunks = _filter_and_rebuild_index(index, existing_chunks, document_id, filename)
                    if remaining_chunks and new_idx is not None:
                        save_index(new_idx, remaining_chunks, index_path, metadata_path)
                    else:
                        # Clean up index files if no chunks remain
                        try:
                            if index_path.exists():
                                index_path.unlink()
                        except Exception:
                            pass
                        try:
                            if metadata_path.exists():
                                metadata_path.unlink()
                        except Exception:
                            pass
                except Exception:
                    pass

        # Also check other subject vector stores in user directory if any
        user_vec_dir = get_user_vector_dir(user_id)
        if user_vec_dir.exists():
            for meta_file in list(user_vec_dir.glob("*.json")):
                if subject and meta_file.stem == subject:
                    continue  # already handled above
                idx_file = meta_file.with_suffix(".index")
                if not idx_file.exists():
                    continue
                try:
                    index, chunks = load_index(idx_file, meta_file)
                    has_doc_chunk = any(
                        c.get("document_id") == document_id or c.get("source") == filename or c.get("filename") == filename
                        for c in chunks
                    )
                    if has_doc_chunk:
                        new_idx, remaining_chunks = _filter_and_rebuild_index(index, chunks, document_id, filename)
                        if remaining_chunks and new_idx is not None:
                            save_index(new_idx, remaining_chunks, idx_file, meta_file)
                        else:
                            try:
                                if idx_file.exists():
                                    idx_file.unlink()
                            except Exception:
                                pass
                            try:
                                if meta_file.exists():
                                    meta_file.unlink()
                            except Exception:
                                pass
                except Exception:
                    pass

        # 3. Delete from database
        delete_user_document(document_id, user_id)

        return True, f"'{filename}' and all associated data were successfully deleted."
    except Exception as e:
        return False, f"Failed to delete document: {str(e)}"


# --------------------------------------------------
# BACKGROUND ASYNCHRONOUS DOCUMENT/IMAGE WORKER
# --------------------------------------------------

def process_document_background(
    job_id: int,
    user_id: int,
    document_id: int,
    file_path: str,
    filename: str,
    subject: str,
    force_ocr: bool = False
):
    """
    Background worker executed in a separate daemon thread.
    Extracts text (with OCR for images/scanned PDFs), chunks content, creates Ollama embeddings,
    builds FAISS index, caches full text in database, and updates persistent job status.
    """
    try:
        # Step 1: Mark processing started (10%)
        update_job_status(job_id, status="processing", progress=10)

        # Step 2: Extract text from file (10% -> 35%)
        def on_extract_progress(current_page, total_pages):
            pct = 10 + int((current_page / max(1, total_pages)) * 25)
            update_job_status(job_id, status="processing", progress=min(35, pct))
            
        pages = extract_document_pages(file_path, force_ocr=force_ocr, progress_callback=on_extract_progress)
        if not pages:
            ext = Path(filename).suffix.lower()
            if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
                raise ValueError("No readable text was detected in this image. Please ensure the image is clear and contains readable text.")
            elif ext in (".docx", ".doc"):
                raise ValueError(f"No extractable text found in Word document: '{filename}'. Please verify the file is not corrupted or password protected.")
            else:
                raise ValueError(f"No readable text or extractable content found in: '{filename}'.")

        # Cache full extracted text in DB for instant cross-feature reuse
        full_text = "\n\n".join(p["text"] for p in pages if p.get("text"))
        save_document_extracted_text(document_id, full_text)

        update_job_status(job_id, status="processing", progress=35)

        # Step 3: Chunk text (55%)
        new_chunks = []
        for page in pages:
            page_chunks = chunk_text(page["text"])
            for chunk in page_chunks:
                new_chunks.append({
                    "text": chunk,
                    "subject": subject,
                    "page": page.get("page", 1),
                    "source": filename,
                    "user_id": user_id,
                    "document_id": document_id
                })

        if not new_chunks:
            raise ValueError("Failed to create text chunks from uploaded material.")
        update_job_status(job_id, status="processing", progress=55)

        # Step 4: Load any existing chunks for this user & subject to merge (65%)
        index_path, metadata_path = get_user_subject_index_paths(user_id, subject)
        _, existing_chunks = load_index(index_path, metadata_path)

        # Remove previous chunks of the same document if re-uploading
        combined_chunks = [c for c in existing_chunks if c.get("source") != filename]
        combined_chunks.extend(new_chunks)

        # Step 5: Compute Ollama embeddings (75% -> 88%)
        update_job_status(job_id, status="processing", progress=75)
        texts = [item["text"] for item in combined_chunks]

        def on_embed_progress(current_count: int, total_count: int):
            pct = 75 + int((current_count / max(1, total_count)) * 13)
            update_job_status(job_id, status="processing", progress=min(88, pct))

        try:
            embeddings = get_embeddings(texts, batch_size=4, progress_callback=on_embed_progress)
        except TypeError:
            embeddings = get_embeddings(texts)

        # Step 6: Build FAISS index and save persistently (90%)
        update_job_status(job_id, status="processing", progress=90)
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)
        index.add(embeddings)

        save_index(index, combined_chunks, index_path, metadata_path)

        # Step 7: Completed (100%)
        update_job_status(job_id, status="completed", progress=100)

    except Exception as e:
        error_msg = str(e)
        update_job_status(job_id, status="failed", progress=0, error_message=error_msg)


# Maintain backward compatibility alias
process_pdf_background = process_document_background


def start_background_indexing(
    user_id: int,
    filename: str,
    file_path: str,
    subject: str,
    file_size: int = 0,
    file_type: str = "pdf",
    force_ocr: bool = False
) -> Tuple[int, int]:
    """
    Creates persistent job & document records and spawns a background thread.
    Returns (document_id, job_id).
    """
    doc_id, job_id = create_document_job(
        user_id=user_id,
        filename=filename,
        file_path=str(file_path),
        subject=subject,
        file_size=file_size,
        file_type=file_type
    )

    worker_thread = threading.Thread(
        target=process_document_background,
        args=(job_id, user_id, doc_id, str(file_path), filename, subject, force_ocr),
        daemon=True,
        name=f"DocWorker-Job-{job_id}"
    )
    worker_thread.start()

    return doc_id, job_id